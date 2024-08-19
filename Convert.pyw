########################################################################################################
# Projet : PDFtoDocx                                                                                   #
# Auteur : Soradev                                                                                     #
# Version : 1.0.0                                                                                      #
########################################################################################################
# Description :                                                                                        #
#   Concert PDF to DOCX                                                                                #
########################################################################################################
# For any questions or contributions, please contact the author at sora.dev.pro@gmail.com              #
########################################################################################################


import fitz
import docx
from docx.shared import Pt, Inches
import os
import logging
from shutil import which
import subprocess


logging.basicConfig(
    level=logging.INFO,
    format="[%(levelname)s] %(message)s"
)

class Converter:
    def __init__(self, pdf_file=None, password=None):
        '''Initialisation du convertisseur avec le fichier PDF.'''
        self.filename_pdf = pdf_file
        self.password = str(password or "")

        if not pdf_file:
            raise ValueError("Un fichier PDF doit être fourni.")

        self._fitz_doc = fitz.Document(pdf_file)

    def close(self):
        '''Fermeture du document PDF.'''
        self._fitz_doc.close()

    def convert_with_libreoffice(self, output_file):
        '''Convertir un PDF en DOCX en utilisant LibreOffice (Linux)'''
        cmd = f"lowriter --convert-to docx '{self.filename_pdf}'"
        subprocess.run(cmd, shell=True, check=True)
        new_file = self.filename_pdf.replace('.pdf', '.docx')
        os.rename(new_file, output_file)

    def convert_with_winword(self, output_file):
        '''Convertir un PDF en DOCX en utilisant Microsoft Word (Windows)'''
        try:
            from win32com import client
            word = client.DispatchEx('Word.Application')
            print(f"Tentative d'ouverture du fichier: {self.filename_pdf}")
            doc = word.Documents.Open(os.path.abspath(self.filename_pdf))
            doc.SaveAs(output_file, FileFormat=16) 
            doc.Close()
            word.Quit()
        except Exception as e:
            logging.error(f"Erreur lors de l'ouverture ou de la conversion du fichier : {e}")

    def convert_with_pymupdf(self, output_file):
        '''Convertir un PDF en DOCX en utilisant PyMuPDF et python-docx'''
        doc = docx.Document()
        for page_num in range(len(self._fitz_doc)):
            page = self._fitz_doc.load_page(page_num)
            text = page.get_text("text")
            doc.add_paragraph(text)
            for img in page.get_images(full=True):
                xref = img[0]
                base_image = self._fitz_doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_name = f'image_{page_num}_{xref}.{base_image["ext"]}'
                
                with open(image_name, "wb") as img_file:
                    img_file.write(image_bytes)
                
                doc.add_picture(image_name, width=Inches(2)) 
                os.remove(image_name)
            doc.add_page_break()
        doc.save(output_file)


    def convert(self, output_file=None):
        '''Effectue la conversion complète du PDF en DOCX'''
        if not output_file:
            output_file = self.filename_pdf.replace(".pdf", ".docx")
        
        if which('libreoffice'):
            logging.info("Conversion avec LibreOffice")
            self.convert_with_libreoffice(output_file)
        elif os.name == 'nt':  
            logging.info("Conversion avec Microsoft Word")
            self.convert_with_winword(output_file)
        else:
            logging.info("Conversion avec PyMuPDF et python-docx")
            self.convert_with_pymupdf(output_file)

        logging.info(f"Conversion terminée avec succès. Fichier enregistré sous {output_file}")

if __name__ == "__main__":
    pdf_path = "chemin/vers/votre_fichier.pdf"  
    output_path = pdf_path.replace(".pdf", ".docx")

    converter = Converter(pdf_file=pdf_path)
    converter.convert(output_file=output_path)
    converter.close()
